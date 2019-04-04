#%%
import pathlib
import os
from docx import Document


#%%
pathlib.Path('certificates').mkdir(parents=True, exist_ok=True) 


#%%
fname = pathlib.Path("names.txt")


#%%
with open(fname) as f:
    name_list = f.readlines()


#%%
for name in name_list:
    name = name.strip()
    document = Document()
    document.add_heading('Certificate of Participation', 0)
    document.add_heading('This acknowledges:', level=1)
    document.add_paragraph()
    second_paragraph = document.add_paragraph()
    first_run = second_paragraph.add_run(f'{name}')
    first_run.bold=True
    document.add_heading('tried the awesome automation examples', level=1)
    document.save(pathlib.Path(f'certificates/Certificate-{name}.docx'))


